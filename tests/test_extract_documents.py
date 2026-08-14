import json
import subprocess
import traceback

import pytest
from docx import Document

from chinese_exam_kit.extract import (
    CapabilityUnavailable,
    PageText,
    UnavailableOcr,
    build_question_candidates,
    extract_document,
    write_extraction_artifacts,
)


def test_markdown_extracts_without_optional_tools(tmp_path):
    source = tmp_path / "exam.md"
    source.write_text("# 一、信息类文本阅读\n\n1. 原创题目", encoding="utf-8")

    result = extract_document(source)

    assert result.mode == "text_layer"
    assert result.pages == (PageText(number=1, text="# 一、信息类文本阅读\n\n1. 原创题目"),)


def test_sparse_pdf_requires_available_ocr(tmp_path):
    source = tmp_path / "scan.pdf"
    source.write_bytes(b"public test fixture")
    rendered = tmp_path / "page-1.png"
    rendered.write_bytes(b"image fixture")

    with pytest.raises(CapabilityUnavailable, match="OCR"):
        extract_document(
            source,
            text_reader=lambda _: [""],
            renderer=lambda *_: [rendered],
            ocr=UnavailableOcr("OCR unavailable; provide an OcrProvider"),
        )


def test_docx_extracts_paragraphs_as_one_page(tmp_path):
    source = tmp_path / "paper.docx"
    document = Document()
    document.add_heading("原创试卷", level=1)
    document.add_paragraph("1. 请概括材料的主要内容。")
    document.save(source)

    result = extract_document(source)

    assert result.mode == "text_layer"
    assert result.pages == (PageText(1, "原创试卷\n1. 请概括材料的主要内容。"),)


def test_docx_keeps_table_text_in_document_order(tmp_path):
    source = tmp_path / "paper.docx"
    document = Document()
    document.add_paragraph("材料说明")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "原创数据"
    document.add_paragraph("1. 分析表格。")
    document.save(source)

    result = extract_document(source)

    assert result.pages[0].text == "材料说明\n指标\t原创数据\n1. 分析表格。"


def test_text_reader_removes_utf8_byte_order_mark(tmp_path):
    source = tmp_path / "paper.txt"
    source.write_bytes("\ufeff原创文本".encode("utf-8"))

    result = extract_document(source)

    assert result.pages[0].text == "原创文本"


class _LocalOcr:
    name = "test-local-ocr"

    def available(self):
        return True

    def recognize(self, image):
        return {"page-2.png": "2. 原创扫描页题目"}[image.name]


def test_mixed_pdf_ocrs_only_sparse_pages_and_preserves_page_numbers(tmp_path):
    source = tmp_path / "mixed.pdf"
    source.write_bytes(b"public fixture")
    images = []
    for number in (1, 2):
        image = tmp_path / f"page-{number}.png"
        image.write_bytes(b"image fixture")
        images.append(image)
    dense = "1. " + "原创文本内容" * 20

    result = extract_document(
        source,
        text_reader=lambda _: [dense, ""],
        renderer=lambda *_: images,
        ocr=_LocalOcr(),
    )

    assert result.mode == "hybrid"
    assert result.pages == (
        PageText(1, dense),
        PageText(2, "2. 原创扫描页题目"),
    )


def test_sparse_pdf_reports_missing_rendered_page(tmp_path):
    source = tmp_path / "scan.pdf"
    source.write_bytes(b"public fixture")
    first_page = tmp_path / "page-1.png"
    first_page.write_bytes(b"image fixture")

    with pytest.raises(ValueError, match="第 2 页"):
        extract_document(
            source,
            text_reader=lambda _: ["", ""],
            renderer=lambda *_: [first_page],
            ocr=_LocalOcr(),
        )


def test_pdf_renderer_failure_does_not_echo_local_source_path(monkeypatch, tmp_path):
    from chinese_exam_kit.extract import documents

    source = tmp_path / "sensitive-name.pdf"
    source.write_bytes(b"public fixture")
    monkeypatch.setattr(documents.shutil, "which", lambda _: "/usr/bin/pdftoppm")

    def fail(command, **kwargs):
        raise subprocess.CalledProcessError(1, command, stderr="local failure")

    monkeypatch.setattr(documents.subprocess, "run", fail)

    with pytest.raises(RuntimeError) as captured:
        extract_document(
            source,
            text_reader=lambda _: [""],
            ocr=_LocalOcr(),
        )

    rendered = "".join(traceback.format_exception(captured.value))
    assert str(tmp_path) not in rendered


def test_unsupported_document_type_is_rejected_before_optional_tools(tmp_path):
    source = tmp_path / "archive.zip"
    source.write_bytes(b"public fixture")

    with pytest.raises(ValueError, match="不支持的文档格式"):
        extract_document(source)


def test_question_candidates_keep_sections_pages_and_sequence_uncertainty():
    pages = (
        PageText(1, "一、信息类文本阅读\n1. 第一题\n2. 第二题"),
        PageText(2, "语言文字运用\n4. 第四题"),
    )

    candidates = build_question_candidates(pages)

    assert [candidate.to_dict() for candidate in candidates] == [
        {
            "number": 1,
            "page_start": 1,
            "section": "reading_1",
            "text": "1. 第一题",
        },
        {
            "number": 2,
            "page_start": 1,
            "section": "reading_1",
            "text": "2. 第二题",
        },
        {
            "number": 4,
            "page_start": 2,
            "section": "language_use",
            "text": "4. 第四题",
            "sequence_warning": "题号从 2 跳至 4",
        },
    ]


def test_extraction_artifacts_are_deterministic_and_contain_no_source_path(tmp_path):
    first_page_text = "一、信息类文本阅读\n1. 原创题目" + "原创内容" * 20
    extraction = extract_document(
        tmp_path / "unused.pdf",
        text_reader=lambda _: [first_page_text, ""],
        renderer=lambda *_: [tmp_path / "page-1.png", tmp_path / "page-2.png"],
        ocr=_LocalOcrForArtifacts(),
    )
    first = tmp_path / "first"
    second = tmp_path / "second"

    write_extraction_artifacts(extraction, first)
    write_extraction_artifacts(extraction, second)

    names = {"question_text.md", "question_index.json", "extraction_review.json"}
    assert {path.name for path in first.iterdir()} == names
    assert {name: (first / name).read_bytes() for name in names} == {
        name: (second / name).read_bytes() for name in names
    }
    assert "<!-- page:1 -->" in (first / "question_text.md").read_text(encoding="utf-8")
    review = json.loads((first / "extraction_review.json").read_text(encoding="utf-8"))
    assert review == {
        "empty_pages": [2],
        "extraction_mode": "hybrid",
        "sequence_warnings": [],
        "unclassified_questions": [],
    }
    combined = b"".join((first / name).read_bytes() for name in sorted(names))
    assert str(tmp_path).encode() not in combined


class _LocalOcrForArtifacts:
    name = "test-local-ocr"

    def available(self):
        return True

    def recognize(self, image):
        return ""
