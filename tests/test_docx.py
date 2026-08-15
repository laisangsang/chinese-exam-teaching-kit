import json
import os
import tempfile
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from chinese_exam_kit.content.docx import (
    Block,
    DEFAULT_STYLE_PATH,
    ListItem,
    build_all,
    build_one,
    parse_markdown,
)


STYLE_PATH = Path("config/docx_style.json")


def _xml(path: Path, member: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read(member).decode("utf-8")


def test_parse_markdown_keeps_editable_block_structure():
    blocks = parse_markdown(
        "# 原创讲评\n\n普通段落含 **重点** 与 `术语`。\n\n"
        "- 先定位证据\n- 再组织答案\n\n1. 审题\n2. 核验\n\n"
        "| 环节 | 任务 |\n| --- | --- |\n| 课前 | 核对材料 |\n"
    )

    assert blocks == [
        Block("heading", text="原创讲评", level=1),
        Block("paragraph", text="普通段落含 **重点** 与 `术语`。"),
        Block(
            "list",
            items=(
                ListItem(level=0, text="先定位证据", marker="-"),
                ListItem(level=0, text="再组织答案", marker="-"),
            ),
        ),
        Block(
            "list",
            items=(
                ListItem(level=0, text="审题", marker="1."),
                ListItem(level=0, text="核验", marker="2."),
            ),
        ),
        Block("table", rows=(("环节", "任务"), ("课前", "核对材料"))),
    ]


def test_build_one_creates_native_structure_and_distinct_evidence_styles(tmp_path):
    source = tmp_path / "guide.md"
    source.write_text(
        "# 原创示例\n\n## 第一题\n\n"
        "普通段落含 **关键词** 与 `限定语`。\n\n"
        "【官方评分参考】答出两个层次。\n\n"
        "【文本推导】依据原创材料。\n\n"
        "【教学拓展】迁移到同类设问。\n\n"
        "- 定位题干\n- 回到文本\n\n"
        "1. 圈画限定\n2. 整合证据\n\n"
        "| 维度 | 证据 |\n| --- | --- |\n| 内容 | 原创事实 |\n",
        encoding="utf-8",
    )
    destination = tmp_path / "guide.docx"

    built = build_one(source, destination, style_path=STYLE_PATH)

    assert built == destination
    doc = Document(destination)
    assert doc.paragraphs[0].text == "原创示例"
    assert doc.paragraphs[0].style.name == "Title"
    assert any(p.text == "第一题" and p.style.name == "Heading 1" for p in doc.paragraphs)
    styles = {p.text.split("】", 1)[0] + "】": p.style.name for p in doc.paragraphs if p.text.startswith("【")}
    assert styles == {
        "【官方评分参考】": "官方评分参考",
        "【文本推导】": "文本推导",
        "【教学拓展】": "教学拓展",
    }
    assert any(run.bold and run.text == "关键词" for p in doc.paragraphs for run in p.runs)
    assert any(run.style and run.style.name == "行内代码" for p in doc.paragraphs for run in p.runs)
    assert any("PAGE" in p._p.xml for section in doc.sections for p in section.footer.paragraphs)
    assert len(doc.tables) == 1

    document_xml = _xml(destination, "word/document.xml")
    numbering_xml = _xml(destination, "word/numbering.xml")
    assert document_xml.count("<w:numPr>") == 4
    assert "<w:numFmt w:val=\"bullet\"" in numbering_xml
    assert "<w:numFmt w:val=\"decimal\"" in numbering_xml
    assert not any(p.text.startswith(("•", "1. ", "- ")) for p in doc.paragraphs)


def test_style_config_is_resolved_into_explicit_a4_ooxml(tmp_path):
    source = tmp_path / "style.md"
    source.write_text(
        "# 教师讲评\n\n## 证据表\n\n| 项目 | 说明 |\n| --- | --- |\n| 一 | 内容较长用于分配列宽 |\n",
        encoding="utf-8",
    )
    destination = tmp_path / "style.docx"
    build_one(source, destination, style_path=STYLE_PATH)
    config = json.loads(STYLE_PATH.read_text(encoding="utf-8"))
    doc = Document(destination)
    section = doc.sections[0]

    assert section.page_width.twips == config["page"]["width_dxa"]
    assert section.page_height.twips == config["page"]["height_dxa"]
    assert section.left_margin.twips == config["page"]["margins_dxa"]["left"]
    assert section.header_distance.twips == config["page"]["header_footer_distance_dxa"]
    normal = doc.styles["Normal"]
    assert normal.font.size.pt == 11
    assert normal._element.rPr.rFonts.get(qn("w:eastAsia")) == "Noto Sans CJK SC"
    assert normal._element.rPr.rFonts.get(qn("w:ascii")) == "Arial"
    assert normal._element.rPr.rFonts.get(qn("w:hint")) == "eastAsia"
    assert normal._element.rPr.find(qn("w:lang")).get(qn("w:eastAsia")) == "zh-CN"
    heading = doc.styles["Heading 1"]
    assert heading._element.rPr.rFonts.get(qn("w:eastAsia")) == "Noto Sans CJK SC"
    assert heading.paragraph_format.keep_with_next
    assert doc.styles["Title"]._element.pPr.find(qn("w:pBdr")) is None

    table_xml = doc.tables[0]._tbl.xml
    widths = config["table"]
    assert f'w:w="{widths["width_dxa"]}"' in table_xml
    assert f'w:w="{widths["indent_dxa"]}"' in table_xml
    grid_widths = [int(node.get(qn("w:w"))) for node in doc.tables[0]._tbl.tblGrid]
    assert sum(grid_widths) == widths["width_dxa"]
    for row in doc.tables[0].rows:
        assert row._tr.trPr.find(qn("w:trHeight")) is None
        for index, cell in enumerate(row.cells):
            assert int(cell._tc.tcPr.tcW.get(qn("w:w"))) == grid_widths[index]


def test_default_style_is_a_packaged_resource_and_matches_public_config(tmp_path):
    source = tmp_path / "default-style.md"
    destination = tmp_path / "default-style.docx"
    source.write_text("# 默认配置\n\n正文。\n", encoding="utf-8")

    build_one(source, destination)

    assert destination.is_file()
    assert json.loads(DEFAULT_STYLE_PATH.read_text(encoding="utf-8")) == json.loads(
        STYLE_PATH.read_text(encoding="utf-8")
    )


def test_build_all_is_deterministic_and_only_builds_direct_markdown_children(tmp_path):
    content = tmp_path / "content"
    output = tmp_path / "output"
    content.mkdir()
    (content / "b.md").write_text("# 乙\n", encoding="utf-8")
    (content / "a.md").write_text("# 甲\n", encoding="utf-8")
    nested = content / "nested"
    nested.mkdir()
    (nested / "hidden.md").write_text("# 不应构建\n", encoding="utf-8")

    outputs = build_all(content, output, style_path=STYLE_PATH)

    assert [path.name for path in outputs] == ["a.docx", "b.docx"]
    assert all(path.is_file() for path in outputs)


def test_ordered_list_keeps_start_and_parent_continuity_across_nested_bullet(tmp_path):
    source = tmp_path / "lists.md"
    source.write_text(
        "# 原创列表\n\n5. 第五项\n6. 第六项\n\n"
        "1. 父项\n  - 子项\n2. 后项\n",
        encoding="utf-8",
    )
    destination = tmp_path / "lists.docx"

    blocks = parse_markdown(source.read_text(encoding="utf-8"))
    assert blocks[1].items[0] == ListItem(level=0, text="第五项", marker="5.")
    assert [item.marker for item in blocks[2].items] == ["1.", "-", "2."]

    build_one(source, destination, style_path=STYLE_PATH)
    document = Document(destination)
    list_paragraphs = [p for p in document.paragraphs if p.text in {"第五项", "第六项", "父项", "子项", "后项"}]
    first_num_id = list_paragraphs[0]._p.pPr.numPr.numId.val
    assert list_paragraphs[1]._p.pPr.numPr.numId.val == first_num_id
    mixed_num_ids = {p._p.pPr.numPr.numId.val for p in list_paragraphs[2:]}
    assert len(mixed_num_ids) == 1
    assert [p._p.pPr.numPr.ilvl.val for p in list_paragraphs[2:]] == [0, 1, 0]
    numbering_xml = _xml(destination, "word/numbering.xml")
    assert '<w:start w:val="5"' in numbering_xml


def test_build_all_redacts_directory_enumeration_errors(tmp_path, monkeypatch):
    content = tmp_path / "private-content"
    output = tmp_path / "output"
    content.mkdir()

    def fail_iterdir(path):
        raise OSError(f"cannot enumerate {tmp_path}/private-content")

    monkeypatch.setattr(Path, "iterdir", fail_iterdir)
    with pytest.raises(OSError, match="could not enumerate content directory") as error:
        build_all(content, output, style_path=STYLE_PATH)

    assert str(tmp_path) not in str(error.value)


def test_build_one_rejects_symlink_input_and_output_without_exposing_absolute_paths(tmp_path):
    real = tmp_path / "real.md"
    real.write_text("# 原创\n", encoding="utf-8")
    source_link = tmp_path / "linked.md"
    try:
        source_link.symlink_to(real)
    except OSError:
        pytest.skip("symlinks unavailable")

    with pytest.raises(ValueError) as source_error:
        build_one(source_link, tmp_path / "out.docx", style_path=STYLE_PATH)
    assert str(tmp_path) not in str(source_error.value)

    destination = tmp_path / "out.docx"
    outside = tmp_path / "outside.docx"
    outside.write_bytes(b"unchanged")
    destination.symlink_to(outside)
    with pytest.raises(ValueError) as output_error:
        build_one(real, destination, style_path=STYLE_PATH)
    assert str(tmp_path) not in str(output_error.value)
    assert outside.read_bytes() == b"unchanged"


def test_atomic_write_failure_preserves_previous_docx_and_removes_partial(tmp_path, monkeypatch):
    source = tmp_path / "guide.md"
    source.write_text("# 新内容\n", encoding="utf-8")
    destination = tmp_path / "guide.docx"
    destination.write_bytes(b"old-word-file")

    def fail_replace(source_path, destination_path):
        raise OSError("simulated private output path")

    monkeypatch.setattr(os, "replace", fail_replace)
    with pytest.raises(OSError) as error:
        build_one(source, destination, style_path=STYLE_PATH)

    assert str(tmp_path) not in str(error.value)
    assert destination.read_bytes() == b"old-word-file"
    assert list(tmp_path.glob(".*.partial.docx")) == []


def test_missing_source_error_is_path_redacted(tmp_path):
    source = tmp_path / "secret-student-name.md"
    with pytest.raises(FileNotFoundError) as error:
        build_one(source, tmp_path / "out.docx", style_path=STYLE_PATH)

    assert str(tmp_path) not in str(error.value)
    assert "secret-student-name.md" in str(error.value)


def test_temporary_output_creation_failure_is_path_redacted(tmp_path, monkeypatch):
    source = tmp_path / "guide.md"
    source.write_text("# 原创\n", encoding="utf-8")

    def fail_mkstemp(*args, **kwargs):
        raise OSError(f"cannot create {tmp_path}/private.partial")

    monkeypatch.setattr(tempfile, "mkstemp", fail_mkstemp)
    with pytest.raises(OSError) as error:
        build_one(source, tmp_path / "guide.docx", style_path=STYLE_PATH)

    assert str(tmp_path) not in str(error.value)


def test_style_config_rejects_symlinked_parent_directory(tmp_path):
    source = tmp_path / "guide.md"
    source.write_text("# 原创\n", encoding="utf-8")
    linked_config_dir = tmp_path / "linked-config"
    try:
        linked_config_dir.symlink_to(STYLE_PATH.resolve().parent, target_is_directory=True)
    except OSError:
        pytest.skip("symlinks unavailable")

    with pytest.raises(ValueError, match="symlink") as error:
        build_one(
            source,
            tmp_path / "guide.docx",
            style_path=linked_config_dir / STYLE_PATH.name,
        )
    assert str(tmp_path) not in str(error.value)
