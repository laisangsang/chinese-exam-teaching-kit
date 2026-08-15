"""Build editable, printable teacher guides from a conservative Markdown dialect."""

from __future__ import annotations

import html
import json
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips


DEFAULT_STYLE_PATH = Path(__file__).resolve().parents[3] / "config" / "docx_style.json"
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_LIST_RE = re.compile(r"^(\s*)([-+*]|\d+[.)])\s+(.+?)\s*$")
_TABLE_RULE_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*[^*\n]+\*(?!\*))")


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0
    ordered: bool = False
    items: tuple[tuple[int, str], ...] = ()
    rows: tuple[tuple[str, ...], ...] = ()


def _strip_comments(text: str) -> str:
    return re.sub(r"<!--.*?(?:-->|\Z)", "", text, flags=re.DOTALL)


def _split_table_row(line: str) -> tuple[str, ...]:
    source = line.strip().removeprefix("|")
    if source.endswith("|") and not source.endswith(r"\|"):
        source = source[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for character in source:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
        elif character == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return tuple(cells)


def parse_markdown(text: str) -> list[Block]:
    """Parse headings, paragraphs, lists and GitHub-style pipe tables."""

    if not isinstance(text, str):
        raise TypeError("Markdown input must be text")
    lines = _strip_comments(text).replace("\r\n", "\n").replace("\r", "\n").splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            value = " ".join(line.strip() for line in paragraph if line.strip())
            if value:
                blocks.append(Block("paragraph", text=value))
            paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush()
            index += 1
            continue
        heading = _HEADING_RE.match(line)
        if heading:
            flush()
            blocks.append(Block("heading", text=heading.group(2), level=len(heading.group(1))))
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and _TABLE_RULE_RE.match(lines[index + 1]):
            flush()
            rows = [_split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip() and "|" in lines[index]:
                rows.append(_split_table_row(lines[index]))
                index += 1
            width = len(rows[0])
            normalized = [tuple((*row[:width], *("" for _ in range(max(0, width - len(row)))))) for row in rows]
            blocks.append(Block("table", rows=tuple(normalized)))
            continue
        list_match = _LIST_RE.match(line)
        if list_match:
            flush()
            ordered = list_match.group(2)[0].isdigit()
            items: list[tuple[int, str]] = []
            while index < len(lines):
                match = _LIST_RE.match(lines[index])
                if match is None or match.group(2)[0].isdigit() != ordered:
                    break
                indent = len(match.group(1).replace("\t", "    "))
                items.append((min(indent // 2, 8), match.group(3)))
                index += 1
            blocks.append(Block("list", ordered=ordered, items=tuple(items)))
            continue
        paragraph.append(line)
        index += 1
    flush()
    return blocks


def _load_style(path: Path) -> Mapping[str, Any]:
    path = Path(path)
    _reject_symlink_chain(path, role="style config")
    if not path.is_file():
        raise ValueError(f"Word style config is unavailable: {path.name}")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError):
        raise ValueError(f"Word style config is unreadable: {path.name}") from None
    if not isinstance(data, dict) or data.get("schema_version") != 1:
        raise ValueError("Word style config schema_version must be 1")
    try:
        if data["preset_name"] != "compact_reference_guide":
            raise ValueError("Word style config must use compact_reference_guide")
        if data["named_override"] != "teacher_guide_a4":
            raise ValueError("Word style config must use teacher_guide_a4")
        for key in ("cjk_body", "cjk_heading", "latin"):
            if not data["fonts"][key] or not all(isinstance(item, str) and item for item in data["fonts"][key]):
                raise ValueError("Word style font chains cannot be empty")
        if data["table"]["width_dxa"] + data["table"]["indent_dxa"] > data["page"]["content_width_dxa"]:
            raise ValueError("Word table geometry exceeds the A4 text area")
    except (KeyError, TypeError):
        raise ValueError("Word style config is incomplete") from None
    return data


def _set_typeface(target, *, cjk: str, latin: str) -> None:
    target.font.name = latin
    element = target._element if hasattr(target, "_element") else target.element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for slot, family in (("ascii", latin), ("hAnsi", latin), ("eastAsia", cjk), ("cs", latin)):
        rfonts.set(qn(f"w:{slot}"), family)
    rfonts.set(qn("w:hint"), "eastAsia")
    language = rpr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        rpr.append(language)
    language.set(qn("w:val"), "en-US")
    language.set(qn("w:eastAsia"), "zh-CN")


def _font_run(run, style: Mapping[str, Any], *, heading: bool = False) -> None:
    fonts = style["fonts"]
    _set_typeface(
        run,
        cjk=fonts["cjk_heading"][0] if heading else fonts["cjk_body"][0],
        latin=fonts["latin"][0],
    )


def _paragraph_style(doc: DocumentObject, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_shading(style, fill: str) -> None:
    ppr = style.element.get_or_add_pPr()
    shading = ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        ppr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def _set_left_border(style, color: str) -> None:
    ppr = style.element.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), color)
    borders.append(left)


def _configure_styles(doc: DocumentObject, token: Mapping[str, Any]) -> None:
    fonts = token["fonts"]
    body = token["body"]
    normal = doc.styles["Normal"]
    _set_typeface(normal, cjk=fonts["cjk_body"][0], latin=fonts["latin"][0])
    normal.font.size = Pt(body["size_pt"])
    normal.font.color.rgb = RGBColor.from_string(token["colors"]["body"])
    normal.paragraph_format.space_before = Pt(body["before_pt"])
    normal.paragraph_format.space_after = Pt(body["after_pt"])
    normal.paragraph_format.line_spacing = body["line_spacing"]
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    _set_typeface(title, cjk=fonts["cjk_heading"][0], latin=fonts["latin"][0])
    title.font.size = Pt(token["title"]["size_pt"])
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(token["title"]["color"])
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(token["title"]["before_pt"])
    title.paragraph_format.space_after = Pt(token["title"]["after_pt"])
    title.paragraph_format.keep_with_next = True
    title_ppr = title.element.get_or_add_pPr()
    inherited_border = title_ppr.find(qn("w:pBdr"))
    if inherited_border is not None:
        title_ppr.remove(inherited_border)

    for name, key in (("Heading 1", "h1"), ("Heading 2", "h2"), ("Heading 3", "h3")):
        style = doc.styles[name]
        values = token["headings"][key]
        _set_typeface(style, cjk=fonts["cjk_heading"][0], latin=fonts["latin"][0])
        style.font.size = Pt(values["size_pt"])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(values["color"])
        style.paragraph_format.space_before = Pt(values["before_pt"])
        style.paragraph_format.space_after = Pt(values["after_pt"])
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    callouts = {
        "官方评分参考": ("official_fill", "official_border", "official_text"),
        "文本推导": ("inference_fill", "inference_border", "inference_text"),
        "教学拓展": ("extension_fill", "extension_border", "extension_text"),
    }
    for name, color_keys in callouts.items():
        style = _paragraph_style(doc, name)
        style.base_style = normal
        _set_typeface(style, cjk=fonts["cjk_body"][0], latin=fonts["latin"][0])
        style.font.size = Pt(body["size_pt"])
        style.font.color.rgb = RGBColor.from_string(token["colors"][color_keys[2]])
        style.paragraph_format.left_indent = Twips(240)
        style.paragraph_format.right_indent = Twips(160)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = body["line_spacing"]
        _set_shading(style, token["colors"][color_keys[0]])
        _set_left_border(style, token["colors"][color_keys[1]])

    code = doc.styles.add_style("行内代码", WD_STYLE_TYPE.CHARACTER)
    _set_typeface(code, cjk=fonts["cjk_body"][0], latin=fonts["latin"][0])
    code.font.size = Pt(9.5)
    code.font.color.rgb = RGBColor.from_string(token["colors"]["code_text"])


def _configure_section(doc: DocumentObject, token: Mapping[str, Any]) -> None:
    page = token["page"]
    section = doc.sections[0]
    section.page_width = Twips(page["width_dxa"])
    section.page_height = Twips(page["height_dxa"])
    section.top_margin = Twips(page["margins_dxa"]["top"])
    section.right_margin = Twips(page["margins_dxa"]["right"])
    section.bottom_margin = Twips(page["margins_dxa"]["bottom"])
    section.left_margin = Twips(page["margins_dxa"]["left"])
    section.header_distance = Twips(page["header_footer_distance_dxa"])
    section.footer_distance = Twips(page["header_footer_distance_dxa"])


def _add_field(paragraph, instruction: str, token: Mapping[str, Any]) -> None:
    run = paragraph.add_run()
    _font_run(run, token)
    for name, value in (("fldChar", "begin"),):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:fldCharType"), value)
        run._r.append(node)
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    run._r.append(instruction_node)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    run._r.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _header_footer(doc: DocumentObject, title: str, token: Mapping[str, Any]) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(2)
    run = header.add_run(title)
    _font_run(run, token, heading=True)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(token["colors"]["muted"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    run = footer.add_run("教师讲评详案 · 第 ")
    _font_run(run, token)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(token["colors"]["muted"])
    _add_field(footer, " PAGE ", token)
    run = footer.add_run(" 页")
    _font_run(run, token)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(token["colors"]["muted"])


def _append_inline(paragraph, text: str, token: Mapping[str, Any], *, heading: bool = False) -> None:
    position = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(html.unescape(text[position:match.start()]))
            _font_run(run, token, heading=heading)
        value = match.group(0)
        run = paragraph.add_run(html.unescape(value[2:-2] if value.startswith("**") else value[1:-1]))
        _font_run(run, token, heading=heading)
        if value.startswith("**"):
            run.bold = True
        elif value.startswith("`"):
            run.style = "行内代码"
        else:
            run.italic = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(html.unescape(text[position:]))
        _font_run(run, token, heading=heading)


def _new_numbering(doc: DocumentObject, ordered: bool, token: Mapping[str, Any]) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in root.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "multilevel")
    abstract.append(multilevel)
    list_token = token["list"]
    for level in range(9):
        level_node = OxmlElement("w:lvl")
        level_node.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        label = OxmlElement("w:lvlText")
        label.set(qn("w:val"), f"%{level + 1}." if ordered else ("•" if level % 2 == 0 else "–"))
        align = OxmlElement("w:lvlJc")
        align.set(qn("w:val"), "left")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(list_token["marker_aligned_dxa"] + level * list_token["level_step_dxa"]))
        tabs.append(tab)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), str(list_token["after_pt"] * 20))
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(list_token["text_indent_dxa"] + level * list_token["level_step_dxa"]))
        indent.set(qn("w:hanging"), str(list_token["hanging_dxa"]))
        ppr.extend((tabs, spacing, indent))
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        for slot, family in (("ascii", token["fonts"]["latin"][0]), ("hAnsi", token["fonts"]["latin"][0]), ("eastAsia", token["fonts"]["cjk_body"][0])):
            rfonts.set(qn(f"w:{slot}"), family)
        rpr.append(rfonts)
        level_node.extend((start, fmt, label, align, ppr, rpr))
        abstract.append(level_node)
    root.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    num.append(reference)
    root.append(num)
    return num_id


def _number_paragraph(paragraph, num_id: int, level: int) -> None:
    numpr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend((ilvl, numid))


def _column_widths(rows: Sequence[Sequence[str]], total: int) -> list[int]:
    column_count = len(rows[0])
    if column_count == 1:
        return [total]
    weights = [max(6, min(48, max(len(row[index]) for row in rows))) for index in range(column_count)]
    minimum = min(760, total // column_count)
    widths = [max(minimum, round(total * weight / sum(weights))) for weight in weights]
    while sum(widths) > total:
        candidate = max(range(column_count), key=lambda index: widths[index] - minimum)
        widths[candidate] -= 1
    while sum(widths) < total:
        widths[max(range(column_count), key=weights.__getitem__)] += 1
    return widths


def _cell_margins(cell, margins: Mapping[str, int]) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    existing = tcpr.find(qn("w:tcMar"))
    if existing is not None:
        tcpr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tcpr.append(tc_mar)


def _add_table(doc: DocumentObject, rows: Sequence[Sequence[str]], token: Mapping[str, Any]) -> None:
    table_token = token["table"]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    widths = _column_widths(rows, table_token["width_dxa"])
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(table_token["width_dxa"]))
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(table_token["indent_dxa"]))
    properties.append(table_indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), table_token["border_color"])
        borders.append(border)
    properties.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_properties.append(cant_split)
        if row_index == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            row_properties.append(header)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcw = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(widths[column_index]))
            _cell_margins(cell, table_token["cell_margins_dxa"])
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["Normal"]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(table_token["cell_after_pt"])
            paragraph.paragraph_format.line_spacing = table_token["cell_line_spacing"]
            _append_inline(paragraph, rows[row_index][column_index], token, heading=row_index == 0)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:fill"), table_token["header_fill"])
                cell._tc.get_or_add_tcPr().append(shading)


def _callout_style(text: str) -> str:
    stripped = text.lstrip()
    for label in ("官方评分参考", "文本推导", "教学拓展"):
        if stripped.startswith(f"【{label}】"):
            return label
    return "Normal"


def _append_blocks(doc: DocumentObject, blocks: Iterable[Block], token: Mapping[str, Any]) -> None:
    first_title = True
    for block in blocks:
        if block.kind == "heading":
            if block.level == 1 and first_title:
                paragraph = doc.add_paragraph(style="Title")
                first_title = False
                word_heading = True
            else:
                word_level = min(max(block.level - 1, 1), 3)
                paragraph = doc.add_paragraph(style=f"Heading {word_level}")
                word_heading = True
            _append_inline(paragraph, block.text, token, heading=word_heading)
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph(style=_callout_style(block.text))
            _append_inline(paragraph, block.text, token)
        elif block.kind == "list":
            num_id = _new_numbering(doc, block.ordered, token)
            for level, item in block.items:
                paragraph = doc.add_paragraph(style=_callout_style(item))
                _number_paragraph(paragraph, num_id, level)
                _append_inline(paragraph, item, token)
        elif block.kind == "table":
            _add_table(doc, block.rows, token)


def _reject_symlink_chain(path: Path, *, role: str) -> None:
    absolute = path.absolute()
    for candidate in (absolute, *absolute.parents):
        if candidate.is_symlink():
            raise ValueError(f"refusing {role} symlink: {path.name}")


def _safe_source(path: Path) -> Path:
    path = Path(path)
    if ".." in path.parts:
        raise ValueError(f"Markdown source escapes its directory: {path.name}")
    _reject_symlink_chain(path, role="input")
    if not path.is_file():
        raise FileNotFoundError(f"Markdown source is unavailable: {path.name}")
    return path


def _safe_destination(path: Path) -> Path:
    path = Path(path)
    if path.suffix.lower() != ".docx" or ".." in path.parts:
        raise ValueError("Word output must be a .docx file without parent traversal")
    _reject_symlink_chain(path, role="output")
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
    except OSError:
        raise OSError(f"could not prepare Word output: {path.name}") from None
    _reject_symlink_chain(path, role="output")
    if path.exists() and not path.is_file():
        raise ValueError(f"Word output is not a file: {path.name}")
    return path


def _new_document(token: Mapping[str, Any]) -> DocumentObject:
    doc = Document()
    _configure_section(doc, token)
    _configure_styles(doc, token)
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    return doc


def build_one(
    source: Path,
    destination: Path,
    *,
    style_path: Path = DEFAULT_STYLE_PATH,
) -> Path:
    """Build one Word guide without exposing caller paths in failures."""

    source = _safe_source(source)
    destination = _safe_destination(destination)
    token = _load_style(style_path)
    try:
        text = source.read_text(encoding="utf-8")
    except (OSError, UnicodeError):
        raise OSError(f"could not read Markdown source: {source.name}") from None
    blocks = parse_markdown(text)
    title = next((block.text for block in blocks if block.kind == "heading" and block.level == 1), source.stem)
    doc = _new_document(token)
    _header_footer(doc, title, token)
    _append_blocks(doc, blocks, token)
    temporary: Path | None = None
    try:
        descriptor, temporary_name = tempfile.mkstemp(
            prefix=f".{destination.name}.", suffix=".partial.docx", dir=destination.parent
        )
        os.close(descriptor)
        temporary = Path(temporary_name)
        doc.save(temporary)
        os.replace(temporary, destination)
    except OSError:
        raise OSError(f"could not write Word output: {destination.name}") from None
    finally:
        if temporary is not None:
            temporary.unlink(missing_ok=True)
    return destination


def build_all(
    content_dir: Path,
    output_dir: Path,
    *,
    style_path: Path = DEFAULT_STYLE_PATH,
) -> list[Path]:
    """Build direct Markdown children in stable filename order."""

    content_dir = Path(content_dir)
    output_dir = Path(output_dir)
    _reject_symlink_chain(content_dir, role="content directory")
    _reject_symlink_chain(output_dir, role="output directory")
    if not content_dir.is_dir():
        raise FileNotFoundError(f"content directory is unavailable: {content_dir.name}")
    sources = sorted(
        (path for path in content_dir.iterdir() if path.suffix.lower() == ".md" and path.is_file() and not path.is_symlink()),
        key=lambda path: path.name,
    )
    if not sources:
        raise FileNotFoundError(f"no Markdown sources found in: {content_dir.name}")
    outputs: list[Path] = []
    for source in sources:
        destination = output_dir / f"{source.stem}.docx"
        outputs.append(build_one(source, destination, style_path=style_path))
    return outputs
